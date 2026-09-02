#!/usr/bin/env python3
"""Generate a PLACEHOLDER report template so the pipeline is runnable today.

This exists only to prove the plumbing. Replace templates/report_template.docx
with your real template — see README.md for the placeholder syntax to add.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "templates" / "report_template.docx"


def main():
    doc = Document()

    title = doc.add_heading("Daily Activity Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("{{ day_of_week }}, {{ report_date }}")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_heading("1. Summary", level=1)
    doc.add_paragraph("{{ intro }}")

    doc.add_heading("2. Personnel", level=1)
    doc.add_paragraph("The following personnel were involved:")
    doc.add_paragraph("{{ names_joined }}")

    doc.add_heading("3. Schedule of Events", level=1)

    # docxtpl deletes the whole row a {%tr %} tag sits in, so the `for` and
    # `endfor` tags each need their own row, with the body row between them.
    table = doc.add_table(rows=4, cols=5)
    table.style = "Table Grid"

    for cell, heading in zip(
        table.rows[0].cells, ["Time", "Event", "Location", "Category", "Attendees"]
    ):
        cell.paragraphs[0].add_run(heading).bold = True

    table.rows[1].cells[0].paragraphs[0].text = "{%tr for e in events %}"

    body_cells = table.rows[2].cells
    body_cells[0].paragraphs[0].text = "{{ e.time }}"
    body_cells[1].paragraphs[0].text = "{{ e.title }}"
    body_cells[2].paragraphs[0].text = "{{ e.location }}"
    body_cells[3].paragraphs[0].text = "{{ e.category }}"
    body_cells[4].paragraphs[0].text = "{{ e.attendees }}"

    table.rows[3].cells[0].paragraphs[0].text = "{%tr endfor %}"

    doc.add_paragraph()
    doc.add_paragraph("Total events recorded: {{ event_count }}")

    doc.add_heading("4. Remarks", level=1)
    doc.add_paragraph("{{ closing }}")

    footer = doc.add_paragraph("Generated {{ generated_at }}")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(8)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote placeholder template to {OUT}")


if __name__ == "__main__":
    main()
